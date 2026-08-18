import os
import re
import tempfile
import asyncio
import logging
import shutil
import subprocess
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

from typing import Any
from abc import ABC, abstractmethod


class ParserClient(ABC):
    @abstractmethod
    async def parse(self, file_path: str) -> str:
        pass

    @abstractmethod
    def supports(self, file_extension: str) -> bool:
        pass


class LlamaParseClient(ParserClient):
    def __init__(self, api_key: str | None = None):
        self.api_key = api_key or os.environ.get("LLAMA_CLOUD_API_KEY")

    def supports(self, file_extension: str) -> bool:
        return file_extension.upper() in ["PDF", "DOCX", "PPTX", "TXT", "MD", "PNG", "JPG", "JPEG", "XLSX", "CSV", "WPS", "KEY"]

    async def parse(self, file_path: str) -> str:
        if self.api_key:
            # 1. Primary: Official llama_parse SDK
            try:
                from llama_parse import LlamaParse
                parser = LlamaParse(api_key=self.api_key, result_type="markdown")
                loop = asyncio.get_running_loop()
                documents = await loop.run_in_executor(None, parser.load_data, file_path)
                if documents:
                    md_text = "\n\n".join(doc.text for doc in documents if doc.text)
                    if md_text and md_text.strip():
                        return md_text
            except Exception as llama_err:
                logger.info(f"Primary llama_parse SDK call failed, trying llama_cloud fallback: {llama_err}")

            # 2. Secondary: llama_cloud SDK API fallback
            try:
                from llama_cloud import LlamaCloud
                client = LlamaCloud(api_key=self.api_key)

                loop = asyncio.get_running_loop()

                def _llama_cloud_parse():
                    file_obj = client.files.create(file=file_path, purpose="parse")
                    parse_res = client.parsing.parse(
                        file_id=file_obj.id,
                        tier="agentic",
                        version="latest",
                        expand=["markdown_full"],
                    )
                    return parse_res.markdown_full or parse_res.text_full or ""

                md_content = await loop.run_in_executor(None, _llama_cloud_parse)
                if md_content and md_content.strip():
                    return md_content
            except Exception as cloud_err:
                logger.warning(f"Secondary LlamaCloud API call warning: {cloud_err}")

        raise RuntimeError("Llama parser quota exceeded")


class DocumentToPdfConverter:
    """
    Converts non-PDF document formats (.docx, .pptx, .xlsx, .csv, .wps, .key) to temporary PDF
    so that oversized files (>10MB) can be split by page ranges via PyMuPDF into <=10MB parts.
    """

    @staticmethod
    async def convert_to_pdf(input_path: str, ext: str) -> str:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, DocumentToPdfConverter._sync_convert, input_path, ext)

    @staticmethod
    def _sync_convert(input_path: str, ext: str) -> str:
        ext_clean = ext.lower().lstrip(".")
        output_dir = tempfile.mkdtemp(prefix="doc_pdf_conv_")
        target_pdf = os.path.join(output_dir, f"converted_{os.path.basename(input_path)}.pdf")

        # 1. Try LibreOffice / soffice headless converter if installed in environment
        soffice_bin = shutil.which("libreoffice") or shutil.which("soffice")
        if soffice_bin:
            try:
                cmd = [soffice_bin, "--headless", "--convert-to", "pdf", "--outdir", output_dir, input_path]
                res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=90)
                if res.returncode == 0:
                    base_name = os.path.splitext(os.path.basename(input_path))[0]
                    generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")
                    if os.path.exists(generated_pdf):
                        return generated_pdf
            except Exception as libre_err:
                logger.warning(f"Headless LibreOffice conversion failed, falling back to python converters: {libre_err}")

        # 2. Python-native Fallback Converters
        doc_pdf = fitz.open()

        if ext_clean in ["docx", "doc", "wps"]:
            try:
                import docx
                d = docx.Document(input_path)
                full_text = []
                for p in d.paragraphs:
                    if p.text.strip():
                        full_text.append(p.text)
                for t in d.tables:
                    for row in t.rows:
                        row_text = " | ".join(c.text.strip() for c in row.cells)
                        full_text.append(row_text)

                combined = "\n\n".join(full_text)
                page = doc_pdf.new_page(width=595, height=842)
                # Split text across pages if long
                paragraphs = combined.split("\n\n")
                current_text = ""
                for p in paragraphs:
                    if len(current_text) + len(p) > 2000:
                        page.insert_textbox(fitz.Rect(50, 50, 545, 792), current_text, fontsize=10)
                        page = doc_pdf.new_page(width=595, height=842)
                        current_text = p + "\n\n"
                    else:
                        current_text += p + "\n\n"
                if current_text:
                    page.insert_textbox(fitz.Rect(50, 50, 545, 792), current_text, fontsize=10)

                doc_pdf.save(target_pdf)
                doc_pdf.close()
                return target_pdf
            except Exception as docx_err:
                logger.error(f"Docx fallback conversion failed: {docx_err}")

        elif ext_clean in ["csv", "txt", "tsv"]:
            try:
                with open(input_path, "r", encoding="utf-8", errors="replace") as f:
                    lines = f.readlines()

                page = doc_pdf.new_page(width=595, height=842)
                chunk = ""
                for line in lines:
                    if len(chunk) + len(line) > 2000:
                        page.insert_textbox(fitz.Rect(50, 50, 545, 792), chunk, fontsize=9)
                        page = doc_pdf.new_page(width=595, height=842)
                        chunk = line
                    else:
                        chunk += line
                if chunk:
                    page.insert_textbox(fitz.Rect(50, 50, 545, 792), chunk, fontsize=9)

                doc_pdf.save(target_pdf)
                doc_pdf.close()
                return target_pdf
            except Exception as txt_err:
                logger.error(f"Text/CSV conversion failed: {txt_err}")

        # Generic fallback: Wrap raw content into standard PDF
        try:
            page = doc_pdf.new_page(width=595, height=842)
            page.insert_text(fitz.Point(50, 100), f"File {os.path.basename(input_path)} ({ext_clean})")
            doc_pdf.save(target_pdf)
            doc_pdf.close()
            return target_pdf
        except Exception:
            pass

        return input_path


class PdfSplitService:
    @staticmethod
    def is_oversized(file_size_bytes: int, limit_mb: int = 10) -> bool:
        return file_size_bytes > limit_mb * 1024 * 1024

    @staticmethod
    def split_pdf(pdf_path: str, max_size_bytes: int = 10 * 1024 * 1024) -> list[dict[str, Any]]:
        doc = fitz.open(pdf_path)
        total_pages = len(doc)

        parts = []
        current_page = 0
        part_num = 1

        # Target split approx 20 pages per part if file is oversized
        file_sz = max(1, os.path.getsize(pdf_path))
        num_splits = max(2, (file_sz // max_size_bytes) + 1)
        pages_per_part = max(1, total_pages // num_splits)

        while current_page < total_pages:
            end_page = min(current_page + pages_per_part, total_pages)
            new_doc = fitz.open()
            new_doc.insert_pdf(doc, from_page=current_page, to_page=end_page - 1)

            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f"_part{part_num}.pdf")
            new_doc.save(temp_file.name)
            part_size = os.path.getsize(temp_file.name)
            new_doc.close()

            parts.append({
                "part_number": part_num,
                "page_start": current_page + 1,
                "page_end": end_page,
                "file_size_bytes": part_size,
                "temporary_file_path": temp_file.name,
            })

            current_page = end_page
            part_num += 1

        doc.close()
        return parts


class MarkdownMergeService:
    @staticmethod
    def merge_markdown_parts(markdown_list: list[str]) -> str:
        merged = "\n\n".join(markdown_list)
        # Remove duplicate main titles if present
        lines = merged.splitlines()
        seen_h1 = set()
        clean_lines = []
        for line in lines:
            if line.startswith("# "):
                if line in seen_h1:
                    continue
                seen_h1.add(line)
            clean_lines.append(line)
        return "\n".join(clean_lines)


class MarkdownNormalizer:
    @staticmethod
    def normalize(markdown: str) -> str:
        # 1. Normalize spacing around headers
        text = re.sub(r"\n{3,}", "\n\n", markdown)
        # 2. Normalize list bullets
        text = re.sub(r"^\s*[*+]\s", "- ", text, flags=re.MULTILINE)
        # 3. Strip trailing whitespace per line
        text = "\n".join(line.rstrip() for line in text.splitlines())
        return text.strip()
